# backend/report_docx.py
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
 
 
BLUE = RGBColor(0x1F, 0x4E, 0x79)
LBLUE = RGBColor(0x2E, 0x75, 0xB6)
 
 
def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via direct XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)
 
 
def build_docx(report_plan: dict, charts: list, output_path: str):
    doc = DocxDocument()
 
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
 
    # ── Cover title ──────────────────────────────────────────
    title = doc.add_heading(report_plan.get('report_title', 'Business Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLUE
 
    doc.add_paragraph()  # spacer
 
    # ── Executive Summary ─────────────────────────────────────
    doc.add_heading('Executive Summary', 1)
    p = doc.add_paragraph(report_plan.get('executive_summary', ''))
    p.paragraph_format.space_after = Pt(12)
 
    # ── Sections ─────────────────────────────────────────────
    chart_iter = iter(charts)
    for i, section in enumerate(report_plan.get('sections', [])):
        doc.add_heading(section.get('section_title', f'Section {i+1}'), 2)
 
        # Narrative paragraph
        p = doc.add_paragraph(section.get('narrative', ''))
        p.paragraph_format.space_after = Pt(10)
 
        # Chart (if exists)
        chart_data = next(chart_iter, None)
        if chart_data and chart_data.get('png_path') and os.path.exists(chart_data['png_path']):
            doc.add_picture(chart_data['png_path'], width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(chart_data.get('insight_caption', ''))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
 
        # Callout box
        callout = section.get('callout', {})
        if callout.get('text'):
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            _set_cell_bg(cell, 'D6E4F0')
            label_para = cell.add_paragraph()
            label_run = label_para.add_run(callout.get('label', 'KEY INSIGHT'))
            label_run.bold = True
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = BLUE
            text_para = cell.add_paragraph(callout.get('text', ''))
            text_para.runs[0].font.size = Pt(10) if text_para.runs else None
            doc.add_paragraph()  # spacer after callout
 
    # ── Conclusion ────────────────────────────────────────────
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(report_plan.get('conclusion', ''))
 
    doc.save(output_path)
    return output_path
